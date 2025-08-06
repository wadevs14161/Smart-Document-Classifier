import aiofiles
from typing import Optional, Tuple
import PyPDF2
from docx import Document as DocxDocument
from io import BytesIO
import logging
import os

logger = logging.getLogger(__name__)

class DocumentProcessingError(Exception):
    """Custom exception for document processing errors"""
    pass

class DocumentProcessor:
    """Handles document text extraction from various file formats with enhanced support"""
    
    # Maximum file size (10MB)
    MAX_FILE_SIZE = 10 * 1024 * 1024
    
    # Supported MIME types and extensions
    SUPPORTED_MIME_TYPES = {
        'text/plain': 'txt',
        'application/pdf': 'pdf',
        'application/vnd.openxmlformats-officedocument.wordprocessingml.document': 'docx',
        'application/msword': 'doc'
    }
    
    SUPPORTED_EXTENSIONS = ['txt', 'pdf', 'docx', 'doc']
    
    @staticmethod
    def validate_file_integrity(file_path: str, expected_type: str) -> Tuple[bool, str]:
        """Validate file integrity and actual type"""
        try:
            if not os.path.exists(file_path):
                return False, "File does not exist"
            
            file_size = os.path.getsize(file_path)
            if file_size == 0:
                return False, "File is empty"
            
            if file_size > DocumentProcessor.MAX_FILE_SIZE:
                return False, f"File too large ({file_size} bytes). Maximum allowed: {DocumentProcessor.MAX_FILE_SIZE} bytes"
            
            # Check actual MIME type using python-magic (if available)
            try:
                import magic
                mime_type = magic.from_file(file_path, mime=True)
                if mime_type not in DocumentProcessor.SUPPORTED_MIME_TYPES:
                    return False, f"Invalid file type detected: {mime_type}"
            except ImportError:
                logger.warning("python-magic not available, skipping MIME type validation")
            
            return True, "File is valid"
            
        except Exception as e:
            return False, f"File validation error: {str(e)}"
    
    @staticmethod
    async def extract_text_from_file(file_path: str, file_type: str) -> Optional[str]:
        """Extract text content from uploaded files with enhanced error handling"""
        try:
            # Validate file first
            is_valid, validation_message = DocumentProcessor.validate_file_integrity(file_path, file_type)
            if not is_valid:
                raise DocumentProcessingError(validation_message)
            
            logger.info(f"Extracting text from {file_type} file: {file_path}")
            
            if file_type.lower() == 'txt':
                return await DocumentProcessor._extract_from_txt(file_path)
            elif file_type.lower() == 'pdf':
                return await DocumentProcessor._extract_from_pdf(file_path)
            elif file_type.lower() in ['docx', 'doc']:
                return await DocumentProcessor._extract_from_docx(file_path)
            else:
                raise DocumentProcessingError(f"Unsupported file type: {file_type}")
                
        except DocumentProcessingError:
            raise
        except Exception as e:
            logger.error(f"Unexpected error extracting text from {file_path}: {str(e)}")
            raise DocumentProcessingError(f"Text extraction failed: {str(e)}")
    
    @staticmethod
    async def _extract_from_txt(file_path: str) -> str:
        """Extract text from plain text files with encoding detection"""
        encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252']
        
        for encoding in encodings:
            try:
                async with aiofiles.open(file_path, 'r', encoding=encoding) as file:
                    content = await file.read()
                    if content.strip():  # Check if content is not just whitespace
                        logger.info(f"Successfully read text file with {encoding} encoding")
                        return content
            except UnicodeDecodeError:
                continue
            except Exception as e:
                logger.error(f"Error reading text file with {encoding}: {str(e)}")
                continue
        
        raise DocumentProcessingError("Unable to read text file - unsupported encoding or corrupted file")
    
    @staticmethod
    async def _extract_from_pdf(file_path: str) -> str:
        """Extract text from PDF files with enhanced PyPDF2 handling"""
        try:
            async with aiofiles.open(file_path, 'rb') as file:
                pdf_content = await file.read()
                
            if len(pdf_content) == 0:
                raise DocumentProcessingError("PDF file is empty")
            
            pdf_reader = PyPDF2.PdfReader(BytesIO(pdf_content))
            
            if len(pdf_reader.pages) == 0:
                raise DocumentProcessingError("PDF has no pages")
            
            text_content = ""
            pages_processed = 0
            total_pages = len(pdf_reader.pages)
            
            logger.info(f"Processing PDF with {total_pages} pages")
            
            for page_num, page in enumerate(pdf_reader.pages):
                try:
                    page_text = page.extract_text()
                    if page_text and page_text.strip():
                        text_content += page_text + "\n"
                        pages_processed += 1
                        logger.debug(f"Extracted text from PDF page {page_num + 1}")
                    else:
                        logger.warning(f"No text extracted from PDF page {page_num + 1}")
                except Exception as e:
                    logger.warning(f"Failed to extract text from PDF page {page_num + 1}: {str(e)}")
                    continue
            
            if not text_content.strip():
                raise DocumentProcessingError("No readable text found in PDF - document may be image-based or corrupted")
            
            # Clean up extracted text
            text_content = text_content.strip()
            # Remove excessive whitespace
            text_content = ' '.join(text_content.split())
            
            logger.info(f"Successfully extracted text from {pages_processed}/{total_pages} PDF pages ({len(text_content)} characters)")
            return text_content
            
        except DocumentProcessingError:
            raise
        except Exception as e:
            error_message = str(e).lower()
            if "incorrect startxref" in error_message or "corrupted" in error_message:
                raise DocumentProcessingError("PDF file appears to be corrupted")
            elif "encrypted" in error_message or "password" in error_message:
                raise DocumentProcessingError("PDF file is encrypted and cannot be processed")
            elif "not a pdf" in error_message:
                raise DocumentProcessingError("File is not a valid PDF format")
            else:
                raise DocumentProcessingError(f"PDF processing error: {str(e)}")
    
    @staticmethod
    async def _extract_from_docx(file_path: str) -> str:
        """Extract text from DOCX files with enhanced handling for tables and formatting"""
        try:
            async with aiofiles.open(file_path, 'rb') as file:
                docx_content = await file.read()
                
            if len(docx_content) == 0:
                raise DocumentProcessingError("DOCX file is empty")
                
            doc = DocxDocument(BytesIO(docx_content))
            
            text_content = ""
            paragraphs_processed = 0
            tables_processed = 0
            
            # Extract text from paragraphs
            logger.debug("Extracting text from paragraphs")
            for paragraph in doc.paragraphs:
                para_text = paragraph.text.strip()
                if para_text:
                    text_content += para_text + "\n"
                    paragraphs_processed += 1
            
            # Extract text from tables
            logger.debug("Extracting text from tables")
            for table in doc.tables:
                table_text = ""
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            row_text.append(cell_text)
                    if row_text:
                        table_text += " | ".join(row_text) + "\n"
                
                if table_text:
                    text_content += "\n" + table_text + "\n"
                    tables_processed += 1
            
            if not text_content.strip():
                raise DocumentProcessingError("No readable text found in DOCX document")
            
            # Clean up extracted text
            text_content = text_content.strip()
            # Normalize whitespace while preserving line breaks
            lines = [' '.join(line.split()) for line in text_content.split('\n')]
            text_content = '\n'.join(line for line in lines if line.strip())
            
            logger.info(f"Successfully extracted text from DOCX document ({paragraphs_processed} paragraphs, {tables_processed} tables, {len(text_content)} characters)")
            return text_content
            
        except DocumentProcessingError:
            raise
        except Exception as e:
            error_message = str(e).lower()
            if "not a zip file" in error_message or "bad zip" in error_message:
                raise DocumentProcessingError("DOCX file appears to be corrupted or not a valid DOCX format")
            elif "no such file" in error_message:
                raise DocumentProcessingError("Invalid DOCX file structure")
            else:
                raise DocumentProcessingError(f"DOCX processing error: {str(e)}")
    
    @staticmethod
    def get_file_type(filename: str) -> str:
        """Get file type from filename"""
        return filename.split('.')[-1].lower() if '.' in filename else 'unknown'
    
    @staticmethod
    def is_supported_file_type(file_type: str) -> bool:
        """Check if file type is supported"""
        supported_types = ['txt', 'pdf', 'docx', 'doc']
        return file_type.lower() in supported_types
    
    @staticmethod
    def get_content_preview(content: str, max_length: int = 200) -> str:
        """Get a preview of the content"""
        if not content:
            return ""
        
        content = content.strip()
        if len(content) <= max_length:
            return content
        
        return content[:max_length] + "..."

    @staticmethod 
    def read_file_content(filepath: str) -> str:
        """
        Synchronous method to read file content based on extension.
        This matches the function from your notebook for compatibility.
        
        Args:
            filepath: Path to the file
            
        Returns:
            Extracted text content
            
        Raises:
            DocumentProcessingError: If file processing fails
        """
        try:
            filename, file_extension = os.path.splitext(filepath)
            file_extension = file_extension.lower()

            if file_extension == '.txt':
                return DocumentProcessor._read_txt_sync(filepath)
            elif file_extension == '.docx':
                return DocumentProcessor._read_docx_sync(filepath)
            elif file_extension == '.pdf':
                return DocumentProcessor._read_pdf_sync(filepath)
            elif file_extension == '.doc':
                raise DocumentProcessingError("Reading .doc files requires additional dependencies")
            else:
                raise DocumentProcessingError(f"Unsupported file type: {file_extension}")
                
        except DocumentProcessingError:
            raise
        except Exception as e:
            raise DocumentProcessingError(f"Error reading file {filepath}: {str(e)}")
    
    @staticmethod
    def _read_txt_sync(filepath: str) -> str:
        """Synchronous text file reader with encoding detection"""
        encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252', 'iso-8859-1']
        
        for encoding in encodings:
            try:
                with open(filepath, 'r', encoding=encoding, errors='ignore') as f:
                    content = f.read()
                    if content.strip():
                        return content
            except (UnicodeDecodeError, OSError):
                continue
        
        raise DocumentProcessingError("Unable to read text file - unsupported encoding or corrupted file")
    
    @staticmethod
    def _read_docx_sync(filepath: str) -> str:
        """Synchronous DOCX file reader"""
        try:
            doc = DocxDocument(filepath)
            content = ""
            
            # Extract paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    content += paragraph.text + "\n"
            
            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        content += " | ".join(row_text) + "\n"
            
            if not content.strip():
                raise DocumentProcessingError("No readable text found in DOCX")
                
            return content.strip()
            
        except Exception as e:
            raise DocumentProcessingError(f"Error reading DOCX file: {str(e)}")
    
    @staticmethod
    def _read_pdf_sync(filepath: str) -> str:
        """Synchronous PDF file reader using PyPDF2"""
        try:
            with open(filepath, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                if len(pdf_reader.pages) == 0:
                    raise DocumentProcessingError("PDF has no pages")
                
                content = ""
                for page in pdf_reader.pages:
                    try:
                        page_text = page.extract_text()
                        if page_text:
                            content += page_text + "\n"
                    except Exception as e:
                        logger.warning(f"Failed to extract text from PDF page: {str(e)}")
                        continue
                
                if not content.strip():
                    raise DocumentProcessingError("No readable text found in PDF")
                    
                return content.strip()
                
        except DocumentProcessingError:
            raise
        except Exception as e:
            raise DocumentProcessingError(f"Error reading PDF file: {str(e)}")
